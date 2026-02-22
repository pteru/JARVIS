#!/usr/bin/env python3
"""JARVIS DOCX Tool — Create, read, edit, and inspect Word documents."""

import argparse
import json
import os
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def resolve_path(filepath):
    """Resolve a file path, handling NFC/NFD Unicode normalization mismatches."""
    p = Path(filepath)
    if p.exists():
        return str(p)
    # Try NFD normalization (common on Linux with files created on macOS/Windows)
    nfd = unicodedata.normalize("NFD", str(p))
    if os.path.exists(nfd):
        return nfd
    # Try NFC
    nfc = unicodedata.normalize("NFC", str(p))
    if os.path.exists(nfc):
        return nfc
    return str(p)  # Return original, let python-docx raise its own error


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def run_to_dict(run, with_styles=False):
    """Convert a Run to a dict."""
    d = {"text": run.text}
    if with_styles:
        if run.bold:
            d["bold"] = True
        if run.italic:
            d["italic"] = True
        if run.underline:
            d["underline"] = True
        if run.font.name:
            d["font"] = run.font.name
        if run.font.size:
            d["size"] = run.font.size.pt
        if run.font.color and run.font.color.rgb:
            d["color"] = str(run.font.color.rgb)
    return d


def paragraph_to_dict(para, with_styles=False):
    """Convert a Paragraph to a dict."""
    d = {
        "type": "paragraph",
        "text": para.text,
        "style": para.style.name if para.style else None,
    }
    if with_styles:
        d["runs"] = [run_to_dict(r, True) for r in para.runs]
        if para.alignment is not None:
            d["alignment"] = str(para.alignment)
    return d


def table_to_dict(table, with_styles=False):
    """Convert a Table to a dict."""
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    d = {"type": "table", "rows": rows}
    if with_styles and table.style:
        d["style"] = table.style.name
    return d


def get_metadata(doc):
    """Extract core properties as a dict."""
    cp = doc.core_properties
    meta = {}
    for attr in ("author", "title", "subject", "keywords", "category",
                 "comments", "last_modified_by", "revision"):
        val = getattr(cp, attr, None)
        if val:
            meta[attr] = str(val)
    for attr in ("created", "modified"):
        val = getattr(cp, attr, None)
        if val:
            meta[attr] = val.isoformat()
    return meta


def iter_block_items(doc):
    """Yield (type, item) for paragraphs and tables in document order."""
    from docx.oxml.ns import qn
    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield "paragraph", next(
                (p for p in doc.paragraphs if p._element is child), None
            )
        elif child.tag == qn("w:tbl"):
            yield "table", next(
                (t for t in doc.tables if t._element is child), None
            )


# ---------------------------------------------------------------------------
# Markdown conversion helpers
# ---------------------------------------------------------------------------

def _heading_level(style_name):
    """Return heading level (1-9) or 0 if not a heading."""
    if not style_name:
        return 0
    m = re.match(r"Heading\s+(\d)", style_name)
    if m:
        return int(m.group(1))
    if style_name.lower() == "title":
        return 1
    if style_name.lower() == "subtitle":
        return 2
    return 0


def _is_list_style(style_name):
    """Check if a style represents a list item."""
    if not style_name:
        return None, False
    name = style_name.lower()
    if "list bullet" in name:
        return "bullet", True
    if "list number" in name:
        return "number", True
    return None, False


def _runs_to_markdown(runs):
    """Convert runs with inline formatting to markdown text."""
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def paragraph_to_markdown(para):
    """Convert a single paragraph to a markdown line."""
    style_name = para.style.name if para.style else ""

    level = _heading_level(style_name)
    if level > 0:
        prefix = "#" * level + " "
        return prefix + _runs_to_markdown(para.runs)

    list_type, is_list = _is_list_style(style_name)
    if is_list:
        text = _runs_to_markdown(para.runs)
        if list_type == "bullet":
            return f"- {text}"
        else:
            return f"1. {text}"

    return _runs_to_markdown(para.runs)


def table_to_markdown(table):
    """Convert a table to a GitHub-flavored markdown table."""
    rows = []
    for row in table.rows:
        rows.append([cell.text.replace("\n", " ") for cell in row.cells])
    if not rows:
        return ""

    lines = []
    # header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        # pad if fewer cells
        padded = row + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------

def cmd_read(args):
    doc = Document(resolve_path(args.file))

    # Collect sections in document order
    sections = []
    for btype, item in iter_block_items(doc):
        if item is None:
            continue
        if btype == "paragraph":
            sections.append(paragraph_to_dict(item, args.with_styles))
        elif btype == "table":
            sections.append(table_to_dict(item, args.with_styles))

    # Apply range filter (paragraph index only, tables counted too)
    if args.range:
        parts = args.range.split(":")
        start = int(parts[0]) if parts[0] else 0
        end = int(parts[1]) if len(parts) > 1 and parts[1] else len(sections)
        sections = sections[start:end]

    if args.format == "json":
        result = {
            "metadata": get_metadata(doc),
            "sections": sections,
        }
        print(json.dumps(result, default=str, indent=2, ensure_ascii=False))

    elif args.format == "markdown":
        md_lines = []
        for sec in sections:
            if sec["type"] == "paragraph":
                # Rebuild paragraph from the document for markdown
                pass
            elif sec["type"] == "table":
                pass

        # Re-iterate for markdown (need original objects)
        doc2 = Document(resolve_path(args.file))
        items = list(iter_block_items(doc2))
        if args.range:
            parts = args.range.split(":")
            start = int(parts[0]) if parts[0] else 0
            end = int(parts[1]) if len(parts) > 1 and parts[1] else len(items)
            items = items[start:end]

        for btype, item in items:
            if item is None:
                continue
            if btype == "paragraph":
                line = paragraph_to_markdown(item)
                md_lines.append(line)
            elif btype == "table":
                md_lines.append("")
                md_lines.append(table_to_markdown(item))
                md_lines.append("")
        print("\n".join(md_lines))

    else:
        # Plain text
        for sec in sections:
            if sec["type"] == "paragraph":
                print(sec["text"])
            elif sec["type"] == "table":
                for row in sec["rows"]:
                    print("\t".join(row))
                print()


def cmd_create(args):
    # Determine input source
    if args.input:
        input_path = Path(args.input)
        content = input_path.read_text(encoding="utf-8")
        is_json = input_path.suffix.lower() == ".json"
        is_md = input_path.suffix.lower() in (".md", ".markdown")
    else:
        content = sys.stdin.read()
        # Auto-detect: try JSON first
        is_json = False
        is_md = False
        stripped = content.strip()
        if stripped.startswith("{") or stripped.startswith("["):
            is_json = True
        else:
            is_md = True

    # Load template or create blank
    if args.template:
        doc = Document(resolve_path(args.template))
        # Clear existing content from template
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
    else:
        doc = Document()

    if is_json:
        _create_from_json(doc, json.loads(content))
    else:
        _create_from_markdown(doc, content)

    doc.save(args.file)
    print(f"Created {args.file}")


def _create_from_json(doc, data):
    """Build document from JSON structure."""
    if isinstance(data, dict):
        # Structured format: {"metadata": ..., "sections": [...]}
        if "metadata" in data:
            meta = data["metadata"]
            cp = doc.core_properties
            for key in ("author", "title", "subject", "keywords",
                        "category", "comments"):
                if key in meta:
                    setattr(cp, key, meta[key])
        sections = data.get("sections", [])
    elif isinstance(data, list):
        sections = data
    else:
        return

    for sec in sections:
        if not isinstance(sec, dict):
            doc.add_paragraph(str(sec))
            continue
        stype = sec.get("type", "paragraph")
        if stype == "paragraph":
            style = sec.get("style", None)
            p = doc.add_paragraph(style=style)
            if "runs" in sec:
                for r in sec["runs"]:
                    run = p.add_run(r.get("text", ""))
                    if r.get("bold"):
                        run.bold = True
                    if r.get("italic"):
                        run.italic = True
                    if r.get("underline"):
                        run.underline = True
                    if r.get("font"):
                        run.font.name = r["font"]
                    if r.get("size"):
                        run.font.size = Pt(r["size"])
            else:
                p.add_run(sec.get("text", ""))
        elif stype == "table":
            rows = sec.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        table.rows[i].cells[j].text = str(cell_text)


def _create_from_markdown(doc, text):
    """Build document from markdown text."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        m = re.match(r"^(#{1,9})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Table detection (lines starting with |)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table_from_md(doc, table_lines)
            continue

        # Bullet list
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            i += 1
            continue

        # Empty line = skip
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — apply inline formatting
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1


def _add_table_from_md(doc, lines):
    """Parse markdown table lines into a docx table."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows (---)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            if j < ncols:
                table.rows[i].cells[j].text = text


def _add_formatted_runs(paragraph, text):
    """Parse inline markdown formatting into runs."""
    # Simple pattern: **bold**, *italic*, ***bold italic***
    pattern = re.compile(r"(\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True
        pos = m.end()
    # Remaining text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def cmd_edit(args):
    doc = Document(resolve_path(args.file))

    if args.replace or args.replace_all:
        old_text, new_text = (args.replace or args.replace_all)
        replace_all = args.replace_all is not None
        count = 0
        for para in doc.paragraphs:
            if old_text in para.text:
                _replace_in_paragraph(para, old_text, new_text)
                count += 1
                if not replace_all:
                    break
        # Also search tables
        if replace_all or count == 0:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if old_text in para.text:
                                _replace_in_paragraph(para, old_text, new_text)
                                count += 1
                                if not replace_all:
                                    break
        print(f"Replaced {count} occurrence(s)")

    if args.append:
        doc.add_paragraph(args.append)
        print("Appended paragraph")

    if args.insert is not None:
        idx, text = args.insert
        idx = int(idx)
        paras = doc.paragraphs
        if idx < len(paras):
            # Insert before paragraph at index
            new_p = paras[idx]._element
            from docx.oxml.ns import qn
            from lxml import etree
            p_elem = doc.element.body.makeelement(qn("w:p"), {})
            new_p.addprevious(p_elem)
            from docx.text.paragraph import Paragraph
            p = Paragraph(p_elem, doc.element.body)
            p.add_run(text)
        else:
            doc.add_paragraph(text)
        print(f"Inserted paragraph at index {idx}")

    if args.delete is not None:
        idx = args.delete
        paras = doc.paragraphs
        if 0 <= idx < len(paras):
            p = paras[idx]._element
            p.getparent().remove(p)
            print(f"Deleted paragraph at index {idx}")
        else:
            print(f"Error: paragraph index {idx} out of range", file=sys.stderr)
            sys.exit(1)

    if args.set_metadata:
        key, value = args.set_metadata
        cp = doc.core_properties
        if hasattr(cp, key):
            setattr(cp, key, value)
            print(f"Set metadata '{key}' = '{value}'")
        else:
            print(f"Error: unknown metadata key '{key}'", file=sys.stderr)
            sys.exit(1)

    doc.save(args.file)
    print(f"Saved {args.file}")


def _replace_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph while preserving formatting of the first run."""
    # Simple case: single run or no runs
    if len(para.runs) <= 1:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text, 1)
                return
    # Multi-run: rebuild full text, find match, apply to concatenated runs
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return
    new_full = full.replace(old_text, new_text, 1)
    # Clear all runs except first, set first run text to full replacement
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""


def cmd_info(args):
    doc = Document(resolve_path(args.file))
    meta = get_metadata(doc)

    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # Count images (inline shapes)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    # Collect styles used
    styles_used = set()
    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)

    # Word count
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)

    # Section count
    section_count = len(doc.sections)

    info = {
        "file": args.file,
        "metadata": meta,
        "paragraphs": para_count,
        "tables": table_count,
        "images": image_count,
        "sections": section_count,
        "word_count": word_count,
        "styles_used": sorted(styles_used),
    }

    if args.format == "json":
        print(json.dumps(info, indent=2, ensure_ascii=False))
    else:
        print(f"File: {args.file}")
        if meta.get("title"):
            print(f"Title: {meta['title']}")
        if meta.get("author"):
            print(f"Author: {meta['author']}")
        if meta.get("created"):
            print(f"Created: {meta['created']}")
        if meta.get("modified"):
            print(f"Modified: {meta['modified']}")
        print(f"Paragraphs: {para_count}")
        print(f"Tables: {table_count}")
        print(f"Images: {image_count}")
        print(f"Sections: {section_count}")
        print(f"Word count: ~{word_count}")
        if styles_used:
            print(f"Styles: {', '.join(sorted(styles_used))}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="JARVIS DOCX Tool")
    sub = parser.add_subparsers(dest="command", required=True)

    # read
    p = sub.add_parser("read", help="Read DOCX content")
    p.add_argument("file")
    p.add_argument("--format", choices=["text", "json", "markdown"],
                   default="text")
    p.add_argument("--with-styles", action="store_true",
                   help="Include font, color, bold/italic metadata (JSON only)")
    p.add_argument("--range", help="Section range e.g. 10:50")

    # create
    p = sub.add_parser("create", help="Create DOCX from input")
    p.add_argument("file")
    p.add_argument("--input", "-i", help="Input file (markdown or JSON)")
    p.add_argument("--template", help="Template DOCX to copy styles from")

    # edit
    p = sub.add_parser("edit", help="Edit DOCX file")
    p.add_argument("file")
    p.add_argument("--replace", nargs=2, metavar=("OLD", "NEW"),
                   help="Replace first occurrence of text")
    p.add_argument("--replace-all", nargs=2, metavar=("OLD", "NEW"),
                   help="Replace all occurrences of text")
    p.add_argument("--append", help="Append a paragraph")
    p.add_argument("--insert", nargs=2, metavar=("INDEX", "TEXT"),
                   help="Insert paragraph at index")
    p.add_argument("--delete", type=int, metavar="INDEX",
                   help="Delete paragraph at index")
    p.add_argument("--set-metadata", nargs=2, metavar=("KEY", "VALUE"),
                   help="Set document metadata field")

    # info
    p = sub.add_parser("info", help="Show DOCX info")
    p.add_argument("file")
    p.add_argument("--format", choices=["text", "json"], default="text")

    args = parser.parse_args()
    {"read": cmd_read, "create": cmd_create, "edit": cmd_edit,
     "info": cmd_info}[args.command](args)


if __name__ == "__main__":
    main()
